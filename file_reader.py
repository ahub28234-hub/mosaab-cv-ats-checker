
"""
Mosaab CV - File Reader Module
===============================
قارئ ملفات السير الذاتية بصيغ PDF, DOCX, TXT
"""

import os
import re
from pathlib import Path
from typing import Optional, Tuple, Dict

class CVFileReader:
    """قارئ ملفات السير الذاتية"""

    SUPPORTED_FORMATS = ['.pdf', '.docx', '.txt', '.doc']

    def __init__(self):
        self.text = ""
        self.metadata = {}

    def read(self, file_path: str) -> Tuple[str, Dict]:
        """
        قراءة ملف السيرة الذاتية

        Args:
            file_path: مسار الملف

        Returns:
            (النص المستخرج, البيانات الوصفية)
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"الملف غير موجود: {file_path}")

        extension = path.suffix.lower()

        if extension == '.pdf':
            return self._read_pdf(file_path)
        elif extension == '.docx':
            return self._read_docx(file_path)
        elif extension == '.txt':
            return self._read_txt(file_path)
        elif extension == '.doc':
            return self._read_doc(file_path)
        else:
            raise ValueError(f"صيغة غير مدعومة: {extension}. الصيغ المدعومة: {', '.join(self.SUPPORTED_FORMATS)}")

    def _read_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """قراءة ملف PDF"""
        metadata = {'format': 'PDF', 'pages': 0, 'is_scanned': False}
        text = ""

        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                    else:
                        metadata['is_scanned'] = True
                        metadata['scanned_pages'] = metadata.get('scanned_pages', 0) + 1

        except ImportError:
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    metadata['pages'] = len(reader.pages)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        else:
                            metadata['is_scanned'] = True
            except ImportError:
                raise ImportError("يرجى تثبيت pdfplumber أو PyPDF2: pip install pdfplumber")

        if metadata.get('is_scanned'):
            metadata['warning'] = "⚠️ يبدو أن الملف يحتوي على صور مسحوبة ضوئياً. ATS لا يستطيع قراءتها!"

        return text, metadata

    def _read_docx(self, file_path: str) -> Tuple[str, Dict]:
        """قراءة ملف DOCX"""
        metadata = {'format': 'DOCX', 'paragraphs': 0}

        try:
            from docx import Document
            doc = Document(file_path)

            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            metadata['paragraphs'] = len(text_parts)
            text = "\n".join(text_parts)

            # فحص الجداول
            tables_count = len(doc.tables)
            if tables_count > 0:
                metadata['tables_count'] = tables_count
                metadata['warning'] = f"⚠️ يحتوي على {tables_count} جدول. تأكد من أن ATS يستطيع قراءتها."

            return text, metadata

        except ImportError:
            raise ImportError("يرجى تثبيت python-docx: pip install python-docx")

    def _read_txt(self, file_path: str) -> Tuple[str, Dict]:
        """قراءة ملف نصي"""
        metadata = {'format': 'TXT'}

        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()

        return text, metadata

    def _read_doc(self, file_path: str) -> Tuple[str, Dict]:
        """قراءة ملف DOC القديم"""
        metadata = {'format': 'DOC', 'warning': 'صيغة DOC قديمة. يُفضل تحويلها إلى DOCX أو PDF.'}

        try:
            # محاولة استخدام antiword أو textract
            import subprocess
            result = subprocess.run(['antiword', file_path], capture_output=True, text=True)
            if result.returncode == 0:
                text = result.stdout
            else:
                text = "[تعذر قراءة ملف DOC. يرجى تحويله إلى PDF أو DOCX]"
        except:
            text = "[تعذر قراءة ملف DOC. يرجى تحويله إلى PDF أو DOCX]"

        return text, metadata

    def get_file_info(self, file_path: str) -> Dict:
        """الحصول على معلومات الملف"""
        path = Path(file_path)
        stat = path.stat()

        return {
            'name': path.name,
            'size_kb': round(stat.st_size / 1024, 2),
            'extension': path.suffix.lower(),
            'modified': stat.st_mtime
        }


if __name__ == '__main__':
    reader = CVFileReader()
    print("✅ CV File Reader loaded successfully!")
